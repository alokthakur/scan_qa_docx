import sys

from docx import Document

instance_record = {}
class MetaType(type):
    def __call__(cls, *args, **kargs):
        try:
            return instance_record[(cls, args)]
        except:
            instance_record[(cls, args)] = super(MetaType, cls).__call__(*args, **kargs)
            return instance_record[(cls, args)]
        
class Subject(object):
    __metaclass__ = MetaType
    def __init__(self, name):
        self.name = name
        self.modules = []
        self.questions = []
        
    def __repr__(self):
        return '<Subject class | {}>'.format(self.name.encode('utf-8'))
    
    def add_question(self, question):
        self.questions.append(question)
        
    def add_module(self, module):
        self.modules.append(module)


class Module(object):
    __metaclass__ = MetaType
    def __init__(self, name):
        self.name = name
        self.questions = []
        
    def __repr__(self):
        return '<Subject class | {}>'.format(self.name)
    
    def add_question(self, question):
        self.questions.append(question)


class Question(object):
    __metaclass__ = MetaType
    def __init__(self, question, options=[]):
        self.question = question
        self.options = options

    def __repr__(self):
        return '<Subject class | {}>'.format(self.question)

    def add_option(self, options):
        self.options = options

def get_module(doc_iter):
    """
    This should be called when scanner is looking for module
    """
    while True:
        paragraph = next(doc_iter)
        if paragraph.text.startswith('Part'):
            return paragraph.text
def get_question(doc_iter, current_module):
    """
    This should be called when scanner is looking for question
    """
    paragraph = next(doc_iter)
    if paragraph.text.startswith('Part') and not paragraph.text.endswith('?'):
        current_module = paragraph.text
        paragraph = next(doc_iter)
    question = ''
    while not paragraph.text.strip().endswith('?'):
        if not paragraph.text.strip():
            paragraph = next(doc_iter)
            continue
        question += paragraph.text
        paragraph = next(doc_iter)
    question += paragraph.text    
    return question, current_module

def get_option(doc_iter):
    """
    This should be called when scanner is looking for options
    """
    paragraph = next(doc_iter)
    option = ''
    while paragraph.text.strip():
        if option:
            option += '\n'
        option += paragraph.text
        paragraph = next(doc_iter)
    return option

if __name__ == '__main__':
    file_name = sys.argv[1]
    doc  = Document(file_name)
    sub_list = []
    sub_list.append(Subject(doc.paragraphs[0].text))
    current_subject_obj = Subject(doc.paragraphs[0].text)
    current_module = None
    current_question = None
    current_option = None
    counter = 1
    doc_iter = iter(doc.paragraphs[1:])
    while True:
        try:
            if current_module is None:
                current_module = get_module(doc_iter)
                current_module_obj = Module(current_module.encode('utf-8'))
                current_subject_obj.add_module(current_module_obj)
            if current_question is None and current_module is not None:
                current_question, current_module = get_question(doc_iter, current_module)
                current_question_obj = Question(current_question.encode('utf-8'))
                current_module_obj = Module(current_module)
                current_module_obj.add_question(current_question_obj)
            if current_option is None and current_question is not None:
                current_option = get_option(doc_iter)
                current_question_obj.add_option(current_option.encode('utf-8'))
            current_question = None
            current_option = None
        except StopIteration:
            break
    # Print all the questions and options
    print current_subject_obj.name
    print "--------------------------"
    for module in current_subject_obj.modules:
        print module.name
        print "-------------------"
        for question in module.questions:
            print question.question
            print '\n'
            print question.options
    x = 10

